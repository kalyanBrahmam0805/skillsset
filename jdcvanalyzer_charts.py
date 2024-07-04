import streamlit as st
import google.generativeai as genai
from pypdf import PdfReader

from matplotlib import pyplot as plt
import numpy as np

import pandas as pd

import sys

import os
import csv
import time

from decimal import *

rows = [['Name', 'Skills Match', 'Tools Match', 'Relevance', 'Experience Match', 'Qualification Match']]

# Function to load PDF content
def load_pdf(file_path):
    """
    Reads the text content from a PDF file and returns it as a single string.
    """
    reader = PdfReader(file_path)
    text = ""
    for page in reader.pages:
        text += page.extract_text()
    return text

# Function to configure the GenerativeModel
def create_model(api_key):
    genai.configure(api_key=api_key)
    generation_config = {
        "temperature": 0,
        "top_p": 0.95,
        "top_k": 64,
        "max_output_tokens": 8192,
        "response_mime_type": "text/plain",
    }
    safety_settings = [
        {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
        {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
        {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
        {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_MEDIUM_AND_ABOVE"},
    ]

    model = genai.GenerativeModel(
        model_name="gemini-1.5-flash",
        safety_settings=safety_settings,
        generation_config=generation_config,
    )
    return model

# Streamlit application
def main():
    st.title("JD V/s CV Analyzer")
    st.write("Analyzes JD V/s CV using Google Generative AI.")

    # File uploader for resume
    uploaded_resume = st.file_uploader("Upload a resume PDF file", type=["pdf"], key="resume")

    if uploaded_resume is not None:
        resume_text = load_pdf(uploaded_resume)
        st.text_area("Extracted Resume Text", resume_text, height=300)

    # File uploader for job description
    uploaded_job_description = st.file_uploader("Upload a job description PDF file", type=["pdf"], key="job_description")

    job_description = ""
    if uploaded_job_description is not None:
        job_description = load_pdf(uploaded_job_description)
        st.text_area("Extracted Job Description Text", job_description, height=300)


    from docx import Document
    doc = Document()
    doc.add_heading('JD_vs_CV_Analyzer:', level=1)
    doc.add_paragraph("Brahmam")
    #directory = 'C:/Users/venka/OneDrive/Desktop/Skills'
    #directory = st.text_input("Enter the folder path")
    #file_path = os.path.join(directory, 'JD_vs_CV_Analyzer.docx')
    #doc.save(file_path)
    #csv_filename = os.path.join(directory, 'cv_analysis.csv')

    # writing to csv file 
    # with open (csv_filename, 'w') as csvfile:
    #     csvwriter = csv.writer (csvfile)
    #     csvwriter.writerows (rows)

    # # create the directory if it doesn't exist
    # if not os.path.exists(directory):
    #     os.makedirs(directory)
    # doc.save(file_path)
    # print(f"Document saved at {file_path}")

    # sys.exit()

    
    #api_key = st.text_input("Enter your Google API Key:", type="password")
    #TODO - Using Brahmam's API key currently
    api_key = "AIzaSyDXFJgEYuudJjkaP26p8yceysDlnrOym7w"

    if st.button("Analyze") and api_key and uploaded_resume and uploaded_job_description:
        model = create_model(api_key)

        prompt = f"""
        Job Description:
        {job_description}

        Candidate's Resume:
        {resume_text}

        Instructions:
        Extract the following details from the resume and compare them against the job description to evaluate the candidate's suitability for the position.

        Details to Extract from Resume:
        1. Candidate's Native Location: Extract the candidate's native city and country.
        2. Skills: List the candidate's key skills.
        3. Tools: Identify the tools and technologies the candidate is proficient with.
        4. Years of Experience: Determine the total years of relevant experience the candidate has.
        5. Projects Mentioned: Summarize the key projects the candidate has worked on.

        Output to be generated:
        1. Skills Match:
           - Matching Skills:
           - Non-Matching Skills:
        2. Tools Match:
           - Matching Tools:
           - Non-Matching Tools:
        3. Experience Match:
           - Candidate has --- years of experience (required: --- years) - experience_matched/not matched
        4. Project Relevance:
           - project_1_relevance in 10 words
           - project_2_relevance in 10 words
        5. Location Suitability:
           - Candidate's Native Location:
           - Job Location:
           - Distance: distance in miles
           - Candidate probability of willing to travel or relocate based on the distance.
        JD V/s Resume Match %:
        Final Recommendation:
        """

        prompt_piechart = f"""
        Instructions:
        Provide a single numerical value in the form of a percentage for each of the analyzed categories as 
        Skills:
        Tools:
        Relevance:
        """

        prompt_chart_values_skill = f"""

        Instructions:

        Provide a single numerical value in the form of a percentage for skills category. Don't give any other explanation. Don't append % symbol.

        """

        prompt_chart_values_tools = """

        Instructions: Provide a single numerical value in the form of a percentage for tools category. Don't give any other explanation. Don't append % symbol.
        
        """
        prompt_chart_values_relevance = """

        Instructions: Provide a single numerical value in the form of a percentage for relevance category. Don't give any other explanation. Don't append % symbol.
        
        """
        prompt_candidate_name = """

        Provide the candidate name

        """

#Subtract the candidate's years of experience from the required years of experience in JD.
        prompt_experience_level = """

        Calculate and provide the percentage of years of experience against the years of experience required according to JD in the range 0 to 100. Don't give a value which is more than 100. 
 
        """

        prompt_qualification_level = """

        Provide the assessment of qualification matching status. If it is matching, return "Matched" or else return "Not Matched" 
 
        """

        chat_session = model.start_chat(
            history=[
                {"role": "user", "parts": ["Analyze the following resume for skills mentioned:"]},
                {"role": "user", "parts": [prompt]},
            ]
        )
        

        # Generate the Response
        response = chat_session.send_message(prompt)
        from docx import Document
        #create new document
        #doc = Document()
        # Add a title
        #doc.add_heading('JD_vs_CV_Analyzer:', level=1)
        # Add the response text
        #doc.add_paragraph(response.text)
        # save the document
        #directory = 'C:/Users/venka/OneDrive/Desktop/Skills'
        #file_path = os.path.join(directory, 'JD_vs_CV_Analyzer.docx')
        #csv_filename = os.path.join(directory, 'cv_analysis.csv')

        # writing to csv file 
        # with open (csv_filename, 'w') as csvfile:

        # #creating a csv writer object 
        #     csvwriter = csv.writer (csvfile)

        # # writing the fields

        # #csvwriter.writerow (fields)

        # # writing the data rows
        #     csvwriter.writerows (rows)


        # create the directory if it doesn't exist
        # if not os.path.exists(directory):
        #     os.makedirs(directory)
        # doc.save(file_path)
        # print(f"Document saved at {file_path}")
        
        # Generate the response/skills identified
        #response = chat_session.send_message(prompt)
        #st.write(response.text)
        #st.write(f"Document saved at {file_path}")

        response_name = chat_session.send_message (prompt_candidate_name) 
        st.write("Candidate Name: ")
        st.write(response_name.text)


        response_cat0 = chat_session.send_message (prompt_chart_values_skill) 
        st.write("Skills Match Percentage: ")
        st.write(response_cat0.text)

        #doc.add paragraph (response cat0.text)

        response_catl = chat_session.send_message (prompt_chart_values_tools) 
        st.write("Tools Match Percentage: ").append(response_cat1.text)
        #st.write(response_catl.text)

        #doc.add_paragraph (response catl.text)

        response_cat2 = chat_session.send_message(prompt_chart_values_relevance) 
        st.write("Overall Relevance Percentage: ").append(response_cat2.text)
        #st.write(response_cat2.text)


        response_experience_level = chat_session.send_message (prompt_experience_level) 

        experience_level = int (Decimal(response_experience_level.text.strip()))
        if(experience_level > 100):
            experience_level = 100

        st.write("Experience Level Percentage: ")
        st.write(str(experience_level))

        
        response_qualification_level = chat_session.send_message (prompt_qualification_level)
        st.write("Qualification Match Status: ")
        st.write(response_qualification_level.text)


        #doc.add_paragraph (response_cat2.text)

        #doc.save(file_path)

        
        
        categories = ['Skills', 'Tools', 'Relevance', 'Experience Match']

        data = [int (Decimal (response_cat0.text.strip())), int (Decimal (response_catl.text.strip())), int (Decimal (response_cat2.text.strip())), experience_level]

        print(data)

        #Creating plot

        fig = plt.figure(figsize=(10,7))

        plt.pie(data, labels=categories)

        #show plot

        st.pyplot (fig)

        fig_skills = plt.figure(figsize=(10,7))

        categories_skills = ['Skills-Match', 'Skills-NotMatch']

        data_skills = [data[0], 100-data[0]]



        explode = (0.1, 0.1)

        fig_skills, ax = plt.subplots()

        ax.pie (data_skills, explode = explode, labels = categories_skills, autopct='%.1f', startangle=90) 

        st.pyplot (fig_skills)

        # candidate_list = [response_name.text.strip(), response_cat0.text.strip(), response_catl.text.strip(), response_cat2.text.strip(), str(experience_level), response_qualification_level.text.strip()] 
        
        # rows.append(candidate_list)

        # with open (csv_filename, 'w') as csvfile:

        # #creating a csv writer object 
        #     csvwriter = csv.writer (csvfile)

        # # writing the data rows

        #     csvwriter.writerows (rows)

        #introduce delay to avoid 429: Resource exhausted error

        time.sleep(5)




if __name__ == "__main__":
    main()
